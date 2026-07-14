"""Resume text extraction and LLM parsing.

Extraction is local (pypdf / python-docx); parsing reuses the shared
call_openai_json engine. Totals and the display summary are computed
deterministically from the parsed roles — LLM arithmetic on month-precision
dates is unreliable, so the model's own numbers are only a fallback.
"""

import io
import logging
import os
import re
from typing import Any, Dict, List, Optional, Tuple

from backend.services.ai_columns import call_openai_json
from backend.services.import_enrichment import (
    months_between,
    parse_profile_date,
    years_from_months,
)

logger = logging.getLogger(__name__)

RESUME_TEXT_MAX_CHARS = int(os.getenv("RESUME_TEXT_MAX_CHARS", "60000"))
RESUME_PARSE_CHAR_LIMIT = int(os.getenv("RESUME_PARSE_CHAR_LIMIT", "24000"))
LOW_TEXT_THRESHOLD = 200

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

_PARSE_SYSTEM_PROMPT = (
    "You are a deterministic resume parsing engine. Extract ONLY what is literally "
    "written in the resume text. Never infer, never guess, never hallucinate. "
    "If a field is absent, return \"\" for strings, [] for arrays, null for numbers. "
    "Do not derive an email address or phone number from a name or company. "
    "Dates must be formatted \"YYYY-MM\", or \"present\" for a current role's end_date, "
    "or \"\" if unknown. Order roles most recent first. "
    "Output ONLY a single JSON object with exactly the keys of the provided schema. "
    "No markdown, no prose, no explanation."
)

_PARSE_SCHEMA = """{
  "first_name": "", "last_name": "", "full_name": "",
  "email": "", "phone": "", "linkedin": "",
  "location": "", "city": "", "headline": "", "summary": "",
  "skills": [], "certifications": [], "languages": [],
  "total_experience_years": null,
  "max_people_managed": null, "years_team_management": null,
  "roles": [
    {"company": "", "title": "", "location": "", "start_date": "YYYY-MM",
     "end_date": "YYYY-MM|present|", "description": "", "is_current": false}
  ],
  "education": [
    {"institution": "", "degree": "", "field": "", "start_date": "", "end_date": ""}
  ],
  "confidence": "high|medium|low",
  "extraction_notes": ""
}"""


def _extract_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages: List[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(file_bytes))
    parts: List[str] = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_plain(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="replace")


def extract_text(file_bytes: bytes, filename: str) -> Tuple[str, str]:
    """Return (text, status) where status is 'ok' or 'low_text'."""
    ext = os.path.splitext(filename or "")[1].lower()
    if ext == ".pdf":
        text = _extract_pdf(file_bytes)
    elif ext == ".docx":
        text = _extract_docx(file_bytes)
    elif ext in {".txt", ".md"}:
        text = _extract_plain(file_bytes)
    else:
        raise ValueError(f"Unsupported resume format: {ext or 'unknown'}")
    text = re.sub(r"\n{3,}", "\n\n", text or "").strip()[:RESUME_TEXT_MAX_CHARS]
    status = "low_text" if len(text.strip()) < LOW_TEXT_THRESHOLD else "ok"
    return text, status


def parse_resume(text: str) -> Dict[str, Any]:
    model = os.getenv("RESUME_PARSE_OPENAI_MODEL", "gpt-4o-mini").strip() or "gpt-4o-mini"
    user_prompt = (
        f"Return a JSON object with exactly this schema:\n{_PARSE_SCHEMA}\n\n"
        f"Resume text:\n{text[:RESUME_PARSE_CHAR_LIMIT]}"
    )
    parsed = call_openai_json(_PARSE_SYSTEM_PROMPT, user_prompt, model=model, use_web=False, temperature=0.0)
    if not isinstance(parsed, dict) or not parsed:
        raise RuntimeError("Resume parse returned no JSON")
    return _normalize_parsed(parsed)


def _normalize_parsed(parsed: Dict[str, Any]) -> Dict[str, Any]:
    def _s(key: str) -> str:
        value = parsed.get(key)
        return str(value).strip() if value is not None and not isinstance(value, (list, dict)) else ""

    def _str_list(key: str) -> List[str]:
        value = parsed.get(key)
        if not isinstance(value, list):
            return []
        return [str(item).strip() for item in value if str(item).strip()][:100]

    def _num(key: str) -> Optional[float]:
        value = parsed.get(key)
        try:
            return float(value) if value is not None else None
        except (TypeError, ValueError):
            return None

    roles: List[Dict[str, Any]] = []
    for role in parsed.get("roles") or []:
        if not isinstance(role, dict):
            continue
        company = str(role.get("company") or "").strip()
        title = str(role.get("title") or "").strip()
        if not company and not title:
            continue
        roles.append(
            {
                "company": company,
                "title": title,
                "location": str(role.get("location") or "").strip(),
                "start_date": str(role.get("start_date") or "").strip(),
                "end_date": str(role.get("end_date") or "").strip(),
                "description": str(role.get("description") or "").strip()[:2000],
                "is_current": bool(role.get("is_current")),
            }
        )

    education: List[Dict[str, Any]] = []
    for entry in parsed.get("education") or []:
        if not isinstance(entry, dict):
            continue
        institution = str(entry.get("institution") or "").strip()
        degree = str(entry.get("degree") or "").strip()
        if not institution and not degree:
            continue
        education.append(
            {
                "institution": institution,
                "degree": degree,
                "field": str(entry.get("field") or "").strip(),
                "start_date": str(entry.get("start_date") or "").strip(),
                "end_date": str(entry.get("end_date") or "").strip(),
            }
        )

    confidence = _s("confidence").lower()
    return {
        "first_name": _s("first_name"),
        "last_name": _s("last_name"),
        "full_name": _s("full_name"),
        "email": _s("email"),
        "phone": _s("phone"),
        "linkedin": _s("linkedin"),
        "location": _s("location"),
        "city": _s("city"),
        "headline": _s("headline"),
        "summary": _s("summary")[:3000],
        "skills": _str_list("skills"),
        "certifications": _str_list("certifications"),
        "languages": _str_list("languages"),
        "total_experience_years": _num("total_experience_years"),
        "max_people_managed": _num("max_people_managed"),
        "years_team_management": _num("years_team_management"),
        "roles": roles[:25],
        "education": education[:10],
        "confidence": confidence if confidence in {"high", "medium", "low"} else "low",
        "extraction_notes": _s("extraction_notes")[:1000],
    }


def recompute_totals(parsed: Dict[str, Any]) -> Dict[str, Any]:
    """Deterministic totals from role dates; the LLM number is only a fallback."""
    role_months: List[int] = []
    for role in parsed.get("roles") or []:
        start = parse_profile_date(role.get("start_date"))
        end_raw = str(role.get("end_date") or "").strip().lower()
        is_current = role.get("is_current") or end_raw in {"present", "current", "now", ""}
        end = parse_profile_date(role.get("end_date"), default_current=is_current)
        months = months_between(start, end)
        if months > 0:
            role_months.append(months)

    if role_months:
        parsed["total_experience_years"] = years_from_months(sum(role_months))
        # Match the enrichment convention: average tenure over past (non-current) roles.
        past = role_months[1:] if len(role_months) > 1 else role_months
        parsed["avg_years_in_company"] = years_from_months(int(sum(past) / max(len(past), 1)))
    else:
        parsed["avg_years_in_company"] = None
    return parsed


def summarize_resume(parsed: Dict[str, Any]) -> str:
    """Deterministic display summary — no second LLM call."""
    parts: List[str] = []
    headline = parsed.get("headline") or parsed.get("summary", "")[:160]
    if headline:
        parts.append(str(headline).strip())
    roles = parsed.get("roles") or []
    if roles:
        latest = roles[0]
        current = f"{latest.get('title', '')} at {latest.get('company', '')}".strip(" at")
        if current:
            parts.append(f"Latest: {current}")
        parts.append(f"{len(roles)} role(s) on resume")
    total = parsed.get("total_experience_years")
    if total:
        parts.append(f"~{total} yrs total experience")
    skills = parsed.get("skills") or []
    if skills:
        parts.append("Skills: " + ", ".join(skills[:10]))
    education = parsed.get("education") or []
    if education:
        latest_edu = education[0]
        edu = " ".join(filter(None, [latest_edu.get("degree"), latest_edu.get("field")])).strip()
        if edu or latest_edu.get("institution"):
            parts.append(f"Education: {edu or ''} {latest_edu.get('institution', '')}".strip())
    return " · ".join(parts)[:1200]
